"""Generate Pattern Claude reference DOCX files.

This script rebuilds the human-facing Word reference artifacts from the current
skill architecture. It intentionally keeps content compact and operator-oriented:
README overview, quick cheat sheet, finance/investment library, and standalone
gold-standard guides for the pivotal analytical deliverables.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

FONT = "Wix Madefor Display"
FONT_BOLD = "Wix Madefor Display SemiBold"
BLUE = "4280F4"
PURPLE = "3A00FD"
NAVY = "0F4761"
LIGHT_BLUE = "D9E2F3"
ALT_BLUE = "F5F8FF"
GREY = "DDDDDD"
BLACK = "000000"
WHITE = "FFFFFF"
GREEN = "375623"
RED = "C00000"
ORANGE = "C55A11"


SKILL_COUNT = "33 skills"
DATE_LABEL = "June 2026"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = GREY) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_margin(cell, margin: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = mar.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            mar.append(el)
        el.set(qn("w:w"), str(margin if edge in ("left", "right") else 80))
        el.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell) -> None:
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def write_cell(cell, text: str, bold: bool = False, color: str = BLACK, size: int = 8) -> None:
    clear_cell(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = FONT_BOLD if bold else FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BOLD if bold else FONT)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = rgb(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_borders(cell)
        set_cell_margin(cell)
        if widths:
            set_cell_width(cell, widths[idx])
        write_cell(cell, h, bold=True, color=WHITE, size=8)

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            set_cell_borders(cell)
            set_cell_margin(cell)
            if widths:
                set_cell_width(cell, widths[c_idx])
            if r_idx % 2 == 0:
                set_cell_shading(cell, ALT_BLUE)
            write_cell(cell, val, size=8)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    return table


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.font.name = FONT_BOLD
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BOLD)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = rgb(BLACK)
        rest = text[len(bold_lead):]
        if rest:
            rr = p.add_run(rest)
            rr.font.name = FONT
            rr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            rr.font.size = Pt(9)
            rr.font.color.rgb = rgb(BLACK)
    else:
        r = p.add_run(text)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(9)
        r.font.color.rgb = rgb(BLACK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(9)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT_BOLD
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BOLD)
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = rgb(BLUE)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT_BOLD
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BOLD)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = rgb(PURPLE)


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, GREY)
    set_cell_margin(cell, 160)
    clear_cell(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.font.name = FONT_BOLD
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BOLD)
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(NAVY)
    rr = p.add_run(text)
    rr.font.name = FONT
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    rr.font.size = Pt(9)
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def setup_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.2)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(9)

    header = sec.header.paragraphs[0]
    header.text = "PATTERN | Claude Skills Library"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = FONT_BOLD
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BOLD)
        run.font.size = Pt(8)
        run.font.color.rgb = rgb(NAVY)

    footer = sec.footer.paragraphs[0]
    footer.text = f"Internal Reference | {DATE_LABEL}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(7)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run(title)
    title_run.font.name = FONT_BOLD
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BOLD)
    title_run.font.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = rgb(NAVY)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.name = FONT
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = rgb(BLACK)
    return doc


def save(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.author = "Pattern"
    doc.core_properties.comments = "Generated by docs/generate_claude_docs.py"
    doc.save(path)
    print(f"saved {path}")


COMMON_TASKS = [
    ["Start or resume a deal workflow", "deal-master"],
    ["Structure a problem or build issue tree", "mckinsey-consultant"],
    ["Track evidence state / belief register", "analytical-operating-system"],
    ["Full market research project", "market-research + gold-standard guide"],
    ["Size a market", "tam-sam-som-calculator"],
    ["Competitive landscape and moat verdict", "market-research -> competitive-moat-assessment"],
    ["Write an IC memo", "ic-memo-pipeline or ic-memo"],
    ["NTB diligence", "ntb-diligence"],
    ["Build a driver / causal tree", "driver-tree"],
    ["Build a financial model", "financial-model-builder"],
    ["GTM diagnostic workbook", "gtm-metrics-analyzer"],
    ["SaaS metric lookup", "finance-metrics-quickref"],
    ["Two-page executive summary", "executive-summary-writer"],
    ["Pattern Word document", "pattern-docx"],
    ["Pattern investment deck", "pattern-investment-pptx"],
    ["QA a Pattern document", "doc-quality-checker"],
]

LAYERS = [
    ["0 - Deal Orchestration", "deal-master", "Inventories existing work, routes deal workflow, initializes belief register."],
    ["1 - Analytical Method", "mckinsey-consultant", "Owns problem framing, MECE issue trees, Pyramid Principle, and analytical method."],
    ["1b - Evidence Control", "analytical-operating-system", "Owns evidence states, belief registers, Bayesian updates, kill triggers, decision posture."],
    ["2 - Research", "market-research, ic-memo, competitive-moat-assessment, tam-sam-som, GTM, finance quickref", "Gathers evidence and creates research or memo deliverable architecture."],
    ["3 - Quality", "writing-style, claim-scrutinizer, red-team, pre-mortem, boundability", "Hardens prose, claims, adverse cases, and failure modes before production."],
    ["4 - Production", "pattern-docx, pattern-investment-pptx, financial-model-builder, diligence-ddr", "Produces branded output files only after quality layers are complete."],
    ["4b - QA", "doc-quality-checker", "Checks brand compliance, formatting, consistency, and remaining draft language."],
]

FINANCE_SKILLS = [
    ["financial-model-builder", "Foundation", "Builds canonical Input Page, Financial Model Template, and Output Tab from source P&L/BS.", "Load first before downstream finance skills."],
    ["ic-memo", "Investment memo", "Creates three-gate IC memo structure with company quality, sector timing, investment attractiveness, kill criteria.", "Use for standalone memo; pipeline for full workflow."],
    ["ic-memo-pipeline", "Workflow", "Runs intake, market research, diligence, draft, quality passes, DOCX output, and QA.", "Best for full deal write-up."],
    ["ntb-diligence", "Diligence", "Tests whether growth is genuine new-to-brand acquisition vs. base recycling.", "Use when customer acquisition quality matters."],
    ["driver-tree", "Causality", "Decomposes thesis into revenue, margin, cost, capital, and competitive drivers.", "Maps NTBs to MOIC levers."],
    ["deal-workbook-builder", "Workbook", "Builds formula-linked Driver Tree, KPI Tree, NTB Registry, and MOIC Bridge.", "Requires source financial model."],
    ["gtm-metrics-analyzer", "GTM", "Calculates GTM metrics across ARR funnel, pipeline, retention, efficiency, productivity, fiscal maturity.", "Use for board prep metrics and ARR diagnostics."],
    ["kpi-tree-builder", "Post-close", "Turns operating targets into causal KPI trees with owners and tracking cadence.", "Post-close or operating cadence."],
    ["finance-metrics-quickref", "Lookup", "Fast definitions, formulas, benchmarks, and red flags for SaaS and PE metrics.", "Use during analysis, not as a full workbook."],
]

MARKET_ARCH = [
    ["Cover and KPI strip", "4-6 figures that frame the category."],
    ["Context and scope", "Boundary, value chain, in-scope, out-of-scope, adjacent markets."],
    ["Executive Summary", "Company/category, product, market, model, thesis, open questions."],
    ["Market Sizing", "Frame comparison, source/scope table, arithmetic checks."],
    ["Customer Segmentation", "Buyer archetypes, JTBD, budget owner, switching friction."],
    ["Competitive Landscape", "Direct competitors, substitutes, platform-native threats, white space."],
    ["Pricing and Economics", "Pricing archetypes, value capture, margin/retention/payback proxies."],
    ["Technology Trends", "Quantified signals, durable vs. eroding layers, horizon."],
    ["Regulatory/Risk", "Jurisdiction, timing, documented scope, implication, do-not-over-claim."],
    ["Moat Analysis", "Moat type, replicability horizon, binding limits, verdict."],
    ["Strategic Implications", "What a CEO, investor, or operator should do differently."],
    ["Appendix", "Methodology, source labels, arithmetic corrections, open questions."],
]

PIVOTAL_GUIDES = [
    ["Market Research Gold Standard", "docs/Market_Research_Gold_Standard_Guide.docx", "Full standalone market research report."],
    ["IC Memo Gold Standard", "docs/IC_Memo_Gold_Standard_Guide.docx", "Full IC memo with gates, NTBs, returns bridge, risks, and open items."],
    ["Competitive Assessment Gold Standard", "docs/Competitive_Assessment_Gold_Standard_Guide.docx", "Competitive arena, competitor map, moat proof, durability, and displacement paths."],
    ["Executive Summary Gold Standard", "docs/Executive_Summary_Gold_Standard_Guide.docx", "Two-page six-section summary for market research, IC memos, diligence, and strategy."],
    ["Strategic Diligence Gold Standard", "docs/Strategic_Diligence_Gold_Standard_Guide.docx", "NTB registry, evidence states, stress tests, kill triggers, and underwriting handoff."],
]

IC_MEMO_ARCH = [
    ["Investment Recommendation", "PROCEED / REPRICE / RESOLVE FIRST / PASS posture and what would change it."],
    ["Executive Summary", "Two-page summary written last using company, product, market, model, thesis, open questions."],
    ["Company Overview", "Business, customers, scale, management, ownership, decision relevance."],
    ["Investment Thesis / NTBs", "4-7 load-bearing beliefs with evidence state, decision impact, and kill trigger."],
    ["Market and Sector Timing", "Why this is a good sector to invest in today, not just a good sector in general."],
    ["Customer, Product, Competitive Position", "Customer need, product workflow role, differentiation, substitutes, moat."],
    ["Financial Profile and Value Creation", "Revenue, margin, cash flow, operating drivers, value creation plan."],
    ["Valuation, Returns, Scenarios", "Entry valuation, return bridge, base/upside/downside cases, sensitivity."],
    ["Risks, Mitigants, Failure Modes", "Specific risks, leading indicators, mitigants, pre-mortem paths."],
    ["Open Items and IC Decision", "Diligence items tied to decision impact and recommendation conditions."],
]

COMPETITIVE_ARCH = [
    ["Executive Verdict", "Defensible position, against whom, for how long, and why."],
    ["Category Boundary", "Direct competitors, substitutes, adjacent platforms, and non-consumption."],
    ["Customer Choice", "Buyer, user, decision-maker, switching threshold, procurement motion."],
    ["Competitor Landscape", "Named competitors segmented by workflow, buyer, price, and strategic posture."],
    ["Substitutes and Platform Threats", "Manual, internal, bundled, and platform-native alternatives."],
    ["Moat Classification", "Network effects, switching costs, scale, proprietary assets, or efficient scale."],
    ["Evidence and Strength Score", "Mechanism, metric, evidence, and 1-5 strength rating."],
    ["Durability and Erosion", "Replicability horizon, erosion vector, and investment-horizon risk."],
    ["Displacement Paths", "How a challenger, bundle, platform, or regulation could erode position."],
    ["Strategic Implications", "Actions for investment, product, GTM, M&A, or diligence."],
]

EXEC_SUMMARY_SPINE = [
    ["Company Overview", "What is this business/category and why does it matter for the decision?"],
    ["Product Offering", "What platform/functionality exists and what differentiation is structural vs. transient?"],
    ["Market Dynamic", "What market forces make the opportunity attractive or unattractive now?"],
    ["Business Model", "How does value become revenue, margin, retention, and scale?"],
    ["Thesis: What You Need To Believe", "3-5 falsifiable beliefs required for the opportunity to work."],
    ["Open Questions", "4-7 diligence gaps that could change the recommendation."],
]

DILIGENCE_ARCH = [
    ["Decision Context", "Screen, pre-LOI, IC prep, or active diligence posture."],
    ["Thesis and Investment Logic", "One-sentence hypothesis connecting asset quality, market timing, and return path."],
    ["Need-to-Believe Register", "4-7 load-bearing beliefs with evidence states and kill triggers."],
    ["Evidence State Assessment", "Confirmed / supported / mixed / weak / missing status by belief."],
    ["Diligence Plan", "Data requests, source owner, format needed, decision impact, priority."],
    ["Stress Tests", "What breaks each NTB and which leading signal shows it."],
    ["Cross-NTB Correlation", "Which assumptions fail together and create compound downside."],
    ["Underwriting Handoff", "Model assumptions affected, base/downside case changes, kill triggers."],
    ["IC Memo Handoff", "Mapping of NTBs, evidence states, risks, and open items to memo sections."],
    ["Next Actions", "Owner, timeline, evidence required, and decision gate."],
]


def build_readme_doc() -> None:
    doc = setup_doc(
        "Claude Skills Library",
        f"{SKILL_COUNT} across 8 groups | 1 orchestration entry point | 5 functional layers | 2 pipelines | {DATE_LABEL}",
    )
    add_callout(
        doc,
        "Source of truth",
        "Root skill folders in this repo are canonical. Packaged plugin copies must be synced from root before publishing.",
    )
    add_h1(doc, "How Skills Fit Together")
    add_para(doc, "Use deal-master for full deal workflows. Use the most specific skill directly for standalone tasks.")
    add_table(doc, ["Layer", "Skills", "What it owns"], LAYERS, [1800, 3300, 5700])
    add_h1(doc, "Auto-Running Skills")
    add_table(
        doc,
        ["Skill", "Runs when"],
        [
            ["writing-style", "Any formal output: memos, reports, investment theses, PPTX narrative text."],
            ["doc-quality-checker", "After any pattern-docx or pattern-investment-pptx file is produced."],
        ],
        [2500, 8300],
    )
    add_h1(doc, "Primary Workflows")
    add_table(
        doc,
        ["Workflow", "When to use", "Output"],
        [
            ["deal-master", "Start/resume full deal intelligence workflow.", "Routed phase plan and evidence register."],
            ["ic-memo-pipeline", "Full IC memo from intake to branded DOCX.", "10-section IC memo, quality passes, DOCX QA."],
            ["market-research-pipeline", "Standalone gold-standard market research report.", "Pattern DOCX market research report."],
        ],
        [2600, 4200, 4000],
    )
    add_h1(doc, "Market Research Standard")
    add_para(
        doc,
        "Full-mode market research must use the standalone gold-standard guide and the skill reference template. The report must include decision artifacts, source labels, market-sizing arithmetic, competitive substitute logic, moat durability, and strategic implications.",
    )
    add_table(doc, ["Report section", "Minimum standard"], MARKET_ARCH, [3300, 7500])
    add_h1(doc, "Pivotal Gold-Standard Guides")
    add_table(doc, ["Guide", "DOCX", "Use"], PIVOTAL_GUIDES, [3000, 4000, 3800])
    add_h1(doc, "Related Files")
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["README.md", "Repo source of truth for deployed skills, layers, and invocation guide."],
            ["CHEATSHEET.md", "Quick-reference task map, layer sequence, workflow pairs, and brand constants."],
            ["docs/market-research-gold-standard-guide.md", "Standalone source guide for the most thorough market research report."],
            ["docs/Market_Research_Gold_Standard_Guide.docx", "Generated Word guide for full standalone market research reports."],
            ["docs/ic-memo-gold-standard-guide.md", "Standalone source guide for the most thorough IC memo."],
            ["docs/competitive-assessment-gold-standard-guide.md", "Standalone source guide for competitive assessment and moat analysis."],
            ["docs/executive-summary-gold-standard-guide.md", "Standalone source guide for two-page executive summaries."],
            ["docs/strategic-diligence-gold-standard-guide.md", "Standalone source guide for NTB diligence and underwriting handoff."],
            ["market-research/references/gold-standard-report-template.md", "Skill reference loaded during Full-mode market research."],
            ["docs/generate_claude_docs.py", "Regenerates the Word reference artifacts."],
        ],
        [4200, 6600],
    )
    save(doc, DOCS / "Claude_Skills_README.docx")


def build_cheatsheet_doc() -> None:
    doc = setup_doc(
        "Claude Skills - Quick Reference",
        f"{SKILL_COUNT} | 2 pipelines | Pattern Investment Team | {DATE_LABEL}",
    )
    add_h1(doc, "What To Invoke")
    add_table(doc, ["Task", "Skill(s) to invoke"], COMMON_TASKS, [4300, 6500])
    add_h1(doc, "Mandatory Layer Sequence")
    add_table(
        doc,
        ["Step", "Layer", "Rule"],
        [
            ["0", "deal-master", "Inventory existing work and route deal workflow."],
            ["1", "mckinsey-consultant", "Frame the problem and govern methodology."],
            ["2", "analytical-operating-system", "Maintain evidence states, belief register, and decision posture."],
            ["3", "research skill", "Gather evidence and build the analytical output."],
            ["4", "writing-style", "Auto-run prose quality and claim discipline before production."],
            ["5", "claim-scrutinizer / red-team / pre-mortem", "Run where claims, investment risk, or adversarial review matter."],
            ["6", "pattern-docx / pattern-investment-pptx", "Produce the branded file."],
            ["7", "doc-quality-checker", "Auto-run the final document QA gate."],
        ],
        [900, 3500, 6400],
    )
    add_h1(doc, "Market Research Full-Mode Checklist")
    add_table(doc, ["Section", "Required artifact"], MARKET_ARCH, [3300, 7500])
    add_h1(doc, "Pivotal Guide Library")
    add_table(doc, ["Guide", "DOCX", "Use"], PIVOTAL_GUIDES, [3000, 4000, 3800])
    add_h1(doc, "Skills That Always Pair")
    add_table(
        doc,
        ["If you invoke", "Also load"],
        [
            ["deal-master", "mckinsey-consultant, analytical-operating-system"],
            ["ic-memo-pipeline", "mckinsey-consultant, analytical-operating-system"],
            ["market-research", "mckinsey-consultant, competitive-moat-assessment, writing-style"],
            ["ic-memo", "ntb-diligence if NTB_MODE=full, driver-tree, executive-summary-writer"],
            ["driver-tree", "boundability for load-bearing nodes"],
            ["pattern-docx", "doc-quality-checker after file delivery"],
        ],
        [3300, 7500],
    )
    add_h1(doc, "Pattern Brand Quick Ref")
    add_table(
        doc,
        ["Element", "Value"],
        [
            ["Font", "Wix Madefor Display / Wix Madefor Display SemiBold"],
            ["Section headers", "#4280F4"],
            ["Subheaders", "#3A00FD"],
            ["Table headers", "#0F4761 fill + #FFFFFF text"],
            ["Body", "#000000"],
            ["Canonical market research guide", "docs/market-research-gold-standard-guide.md and market-research/references/gold-standard-report-template.md"],
        ],
        [3300, 7500],
    )
    save(doc, DOCS / "Claude_Skills_CheatSheet_v2.docx")


def build_finance_doc() -> None:
    doc = setup_doc(
        "Claude Skill Library - Finance and Investment",
        f"{SKILL_COUNT} | Skill architecture | Finance dependencies | IC memo workflow | Market research handoff | {DATE_LABEL}",
    )
    add_callout(
        doc,
        "Foundational rule",
        "Load financial-model-builder first before downstream finance work when a source P&L, balance sheet, or operating model exists.",
    )
    add_h1(doc, "Contents")
    add_table(
        doc,
        ["#", "Section"],
        [
            ["1", "Finance Layer Map"],
            ["2", "Finance Skill Inventory"],
            ["3", "Hard Codependencies"],
            ["4", "Full IC Memo Sequential Prompts"],
            ["5", "Input Requirements"],
            ["6", "Market Research Handoff To Finance"],
            ["7", "Workflow Chains"],
            ["8", "Trigger Phrases"],
            ["9", "Common Mistakes"],
        ],
        [900, 9900],
    )
    add_h1(doc, "Finance Layer Map")
    add_table(doc, ["Layer", "Skills", "What it owns"], LAYERS, [1800, 3300, 5700])
    add_h1(doc, "Finance Skill Inventory")
    add_table(doc, ["Skill", "Role", "What it does", "Use note"], FINANCE_SKILLS, [2300, 1700, 4500, 2300])
    add_h1(doc, "Hard Codependencies")
    add_table(
        doc,
        ["Skill", "Must load with", "Why"],
        [
            ["financial-model-builder", "Source P&L/BS or operating model context", "Defines the canonical 3-tab model that downstream finance work references."],
            ["ic-memo-pipeline", "mckinsey-consultant, analytical-operating-system", "Methodology and belief-register discipline must be active before memo drafting."],
            ["ic-memo", "market-research, executive-summary-writer", "Market evidence and two-page summary spine are inputs to an IC-ready memo."],
            ["ntb-diligence", "driver-tree, boundability", "NTBs need to map to value drivers and be tested for durability boundaries."],
            ["gtm-metrics-analyzer", "financial-model-builder when a model exists", "Revenue, GP, S&M, and operating metrics should tie to the model instead of being re-entered."],
            ["pattern-docx", "writing-style, claim-scrutinizer, doc-quality-checker", "Formal files require prose hardening, claim integrity, and final QA."],
        ],
        [2400, 3600, 4800],
    )
    add_h1(doc, "IC Memo Workflow")
    add_table(
        doc,
        ["Phase", "Primary skill", "Output gate"],
        [
            ["1 - Intake", "ic-memo / ic-memo-pipeline", "Decision context, thesis, materials inventory, known concerns."],
            ["2 - Market work", "market-research", "Market, customer, competitive, moat, and Gate 2 evidence base."],
            ["3 - Deep diligence", "ntb-diligence, driver-tree, financial-model-builder, GTM, KPI tree", "NTB registry, driver tree, model, GTM workbook, operating KPI plan as needed."],
            ["4 - Draft memo", "ic-memo", "10-section investment memo with executive summary last."],
            ["5 - Quality passes", "writing-style, claim-scrutinizer, red-team, pre-mortem, boundability", "Hardened thesis, risks, assumptions, and underwriting actions."],
            ["6 - Production and QA", "pattern-docx, doc-quality-checker", "Pattern DOCX with zero critical QA issues."],
        ],
        [2100, 3100, 5600],
    )
    add_h1(doc, "Full IC Memo Sequential Prompts")
    add_para(doc, "Use these as copy-paste beats when running a finance or investment memo outside the full pipeline.")
    add_table(
        doc,
        ["Beat", "Prompt pattern", "Expected output"],
        [
            ["1 - Intake", "Write the IC memo for [Company]. Context: [sector, deal type, stage, materials].", "Decision context, source inventory, gating questions."],
            ["2 - Market", "Run market research in IC Memo Mode for [Company/category]. Focus on Gate 2 sector timing.", "Market, customer, competitive, and moat evidence base."],
            ["3 - Diligence", "Run NTB diligence and build the driver tree for the investment thesis.", "NTB registry, driver tree, value-creation levers."],
            ["4 - Model", "Build or audit the financial model from the source P&L/BS.", "Input Page, FMT, Output Tab, variance/margin view."],
            ["5 - Draft", "Draft the 10-section IC memo using the evidence base and model outputs.", "Complete memo draft with executive summary written last."],
            ["6 - Harden", "Run writing-style, claim-scrutinizer, red-team, pre-mortem, and boundability.", "Redlined, hardened memo with kill criteria and open issues."],
            ["7 - Produce", "Generate Pattern DOCX and run doc-quality-checker.", "Final branded IC memo with QA gate cleared."],
        ],
        [1600, 5200, 4000],
    )
    add_h1(doc, "Input Requirements")
    add_table(
        doc,
        ["Workflow", "Inputs that most improve output quality"],
        [
            ["Financial model", "Source workbook, P&L tab, BS tab if available, last actuals period, forecast horizon, units, company name."],
            ["IC memo", "Company, sector, deal type, thesis, entry valuation, hold period, materials path, known IC concerns."],
            ["NTB diligence", "Customer-level revenue, cohort data, gross/net retention, channel history, customer acquisition definitions."],
            ["GTM metrics", "ARR funnel, CRM exports, bookings, pipeline, S&M spend, rep roster, retention tables, board metrics."],
            ["Market research", "Decision to support, geography, time horizon, known sources, source constraints, output format."],
        ],
        [2800, 8000],
    )
    add_h1(doc, "Market Research Handoff To Finance")
    add_para(
        doc,
        "Standalone market research becomes finance-ready only when the market, customer, competitive, economics, moat, and risk sections translate into explicit underwriting assumptions.",
    )
    add_table(
        doc,
        ["Research output", "Finance use"],
        [
            ["Market sizing", "Bounds TAM/SAM/SOM, growth case, penetration ceiling, and market timing risk."],
            ["Customer segmentation", "Defines ICP, ACV, sales motion, retention risk, and expansion assumptions."],
            ["Competitive landscape", "Informs win rate, pricing pressure, displacement path, and exit buyer universe."],
            ["Pricing and economics", "Supports gross margin, CAC payback, LTV/CAC, and unit economics checks."],
            ["Moat and regulation", "Sets multiple durability, downside risk, and kill criteria."],
        ],
        [3300, 7500],
    )
    add_h1(doc, "Workflow Chains")
    add_table(
        doc,
        ["Need", "Recommended chain"],
        [
            ["Full IC memo", "deal-master -> ic-memo-pipeline -> pattern-docx -> doc-quality-checker"],
            ["Standalone finance model", "financial-model-builder -> model checks -> output workbook"],
            ["GTM diagnostic", "financial-model-builder if available -> gtm-metrics-analyzer -> workbook QA"],
            ["Market-led investment thesis", "market-research -> moat assessment -> executive-summary-writer -> ic-memo"],
            ["Downside case", "pre-mortem -> red-team -> boundability -> claim-scrutinizer"],
            ["Post-close operating plan", "kpi-tree-builder -> driver-tree -> 100-day plan"],
        ],
        [3000, 7800],
    )
    add_h1(doc, "Trigger Phrases")
    add_table(
        doc,
        ["Trigger phrase", "Skill"],
        [
            ["build a financial model / 3-tab model / 6+6 analysis", "financial-model-builder"],
            ["write the IC memo / deal memo / investment committee memo", "ic-memo or ic-memo-pipeline"],
            ["run diligence / NTB analysis / is growth real", "ntb-diligence"],
            ["decompose this thesis / what drives MOIC", "driver-tree"],
            ["build a GTM workbook / calculate NDR / board prep metrics", "gtm-metrics-analyzer"],
            ["what is the formula for / define this metric", "finance-metrics-quickref"],
            ["what should we track post-close", "kpi-tree-builder"],
        ],
        [5200, 5600],
    )
    add_h1(doc, "Common Mistakes")
    add_bullets(
        doc,
        [
            "Building a financial model before confirming the source P&L and actual/forecast split.",
            "Using market research as narrative support without translating it into underwriting assumptions.",
            "Treating vendor ROI claims as independent proof of customer economics.",
            "Skipping claim-scrutinizer before producing an IC memo document.",
            "Building post-close KPI trees from pre-diligence assumptions without labeling uncertainty.",
        ],
    )
    add_h1(doc, "Reference Files")
    add_table(
        doc,
        ["Reference", "Use"],
        [
            ["financial-model-builder/references/blueprint.md", "Canonical row-by-row model spec."],
            ["market-research/references/gold-standard-report-template.md", "Full-mode market research report architecture."],
            ["docs/market-research-gold-standard-guide.md", "Human-readable guide for the most thorough market research document."],
            ["mckinsey-consultant/references/investment-evaluation-framework.md", "Investment evaluation gates and diligence questions."],
            ["mckinsey-consultant/references/VALIDATION_FRAMEWORKS.md", "Source validation, CRAAP scoring, and triangulation."],
        ],
        [5000, 5800],
    )
    save(doc, DOCS / "Claude_Skill_Library_External (Finance)_v6.docx")


def build_market_research_doc() -> None:
    doc = setup_doc(
        "Market Research Gold Standard Guide",
        f"How to create the most thorough Pattern market research document | {DATE_LABEL}",
    )
    add_callout(
        doc,
        "BLUF",
        "A thorough market research report is a decision document: it defines the market boundary, proves buyer behavior, reconciles market size, maps competition and substitutes, tests moat durability, and converts findings into underwriting or operating actions.",
    )
    add_h1(doc, "Decision Standard")
    add_bullets(
        doc,
        [
            "State the decision before research begins.",
            "Use a MECE hypothesis tree to identify what must be true.",
            "Separate facts, estimates, hypotheses, vendor claims, and Pattern analytics.",
            "Build artifacts in every major section; prose-only sections are incomplete.",
            "End with actions, underwriting implications, and open questions.",
        ],
    )
    add_h1(doc, "Canonical Report Architecture")
    add_table(doc, ["Section", "Minimum standard"], MARKET_ARCH, [3300, 7500])
    add_h1(doc, "Workflow")
    add_table(
        doc,
        ["Step", "Action", "Output"],
        [
            ["1", "Define the decision and working hypothesis.", "Decision statement, hypothesis, and disconfirming conditions."],
            ["2", "Complete the research brief.", "Issue tree, source plan, success criteria, kill criteria."],
            ["3", "Collect evidence by pyramid level.", "L4 market, L3 customer, L2 competitive, L1 position evidence base."],
            ["4", "Build decision artifacts.", "KPI strip, sizing tables, buyer maps, competitor maps, economics tables, risk tables."],
            ["5", "Check arithmetic and source labels.", "CAGR ties, source/scope reconciliation, evidence labels."],
            ["6", "Draft and harden.", "writing-style, claim-scrutinizer, red-team, then Pattern DOCX and QA."],
        ],
        [900, 4300, 5600],
    )
    add_h1(doc, "Market Sizing Rules")
    add_bullets(
        doc,
        [
            "Separate reference markets, analytic slices, and upside scenarios.",
            "Do not sum overlapping markets unless overlap is removed.",
            "Tie every CAGR to start value, end value, and number of years.",
            "Label Pattern-built estimates separately from third-party market figures.",
            "Treat vendor or commissioned studies as directional unless independently triangulated.",
        ],
    )
    add_para(doc, "Core formula: CAGR = (End value / Start value) ^ (1 / Years) - 1")
    add_h1(doc, "Evidence Labels")
    add_table(
        doc,
        ["Label", "Use for", "Treatment"],
        [
            ["Fact", "Official, primary, or directly sourced evidence.", "Strongest support."],
            ["Estimate", "Reasoned calculation from credible inputs.", "Show assumptions."],
            ["Hypothesis", "Plausible but unproven interpretation.", "State confirming evidence needed."],
            ["Vendor claim", "Vendor, PR, or commissioned customer claim.", "Directional only."],
            ["Pattern analytic", "Internal model, bridge, or scenario.", "Keep separate from consensus."],
        ],
        [1700, 5200, 3900],
    )
    add_h1(doc, "Executive Summary Spine")
    add_table(
        doc,
        ["Section", "What it must answer"],
        [
            ["Company Overview", "What is the company/category and why does it matter?"],
            ["Product Offering", "What platform/functionality is offered, and what differentiation is structural vs. transient?"],
            ["Market Dynamic", "What market forces make this attractive or unattractive now?"],
            ["Business Model", "How does value translate into revenue, margin, retention, and scalability?"],
            ["Thesis", "What must be true to believe the opportunity is attractive?"],
            ["Open Questions", "Which diligence items could change the recommendation?"],
        ],
        [3000, 7800],
    )
    add_h1(doc, "Release Checklist")
    add_bullets(
        doc,
        [
            "Every headline is an insight, not a topic label.",
            "Every table has a what-this-proves interpretation.",
            "Every market size states scope, source type, year, and arithmetic.",
            "Every vendor claim is labeled as vendor or commissioned evidence.",
            "Every moat claim states structural, conditional, transient, or unproven durability.",
            "Every strategic implication changes an action, underwriting assumption, or diligence question.",
        ],
    )
    add_h1(doc, "Paste-Ready Prompt")
    add_callout(
        doc,
        "Prompt",
        "Run a full standalone market research report on [market/company/category]. Decision: [enter/invest/acquire/partner/prioritize/avoid]. Geography: [scope]. Time horizon: [horizon]. Use market-research in Full mode, load mckinsey-consultant first, use the gold-standard architecture, label evidence, show market-sizing arithmetic, include artifacts in every section, then run writing-style, claim-scrutinizer, red-team, pattern-docx, and doc-quality-checker.",
        fill="F5F8FF",
    )
    save(doc, DOCS / "Market_Research_Gold_Standard_Guide.docx")


def build_ic_memo_doc() -> None:
    doc = setup_doc(
        "IC Memo Gold Standard Guide",
        f"How to create the most thorough Pattern investment committee memo | {DATE_LABEL}",
    )
    add_callout(
        doc,
        "BLUF",
        "A full IC memo is a decision document: it recommends proceed, reprice, resolve first, or pass; proves the three investment gates; ties NTBs to evidence; and connects operating drivers to returns.",
    )
    add_h1(doc, "Decision Standard")
    add_bullets(
        doc,
        [
            "Answer Gate 1: why this is a good company.",
            "Answer Gate 2: why this is a good sector to invest in today.",
            "Answer Gate 3: why this is a good investment at this price and structure.",
            "State the case not to do the deal before final recommendation hardening.",
            "Make kill criteria explicit and measurable.",
        ],
    )
    add_h1(doc, "Canonical 10-Section Memo")
    add_table(doc, ["Section", "Minimum standard"], IC_MEMO_ARCH, [3300, 7500])
    add_h1(doc, "Need-To-Believe Register")
    add_table(
        doc,
        ["NTB field", "What it must capture"],
        [
            ["NTB", "A load-bearing, falsifiable belief required for the thesis to work."],
            ["Current evidence", "Named evidence and source quality."],
            ["Evidence state", "Confirmed, supported, mixed, weak, or missing."],
            ["Decision impact", "Proceed, reprice, resolve first, or pass implication."],
            ["Confirming data", "Specific data needed to prove the belief."],
            ["Kill trigger", "Finding that changes the recommendation."],
        ],
        [3000, 7800],
    )
    add_h1(doc, "Returns Bridge")
    add_table(
        doc,
        ["Component", "Question"],
        [
            ["EBITDA growth", "What revenue and margin drivers create earnings growth?"],
            ["Margin improvement", "What operational levers expand EBITDA margin?"],
            ["Multiple expansion/contraction", "What exit multiple is justified and why?"],
            ["Leverage paydown / cash", "How much value comes from deleveraging or cash generation?"],
        ],
        [3300, 7500],
    )
    add_h1(doc, "Quality Stack")
    add_bullets(
        doc,
        [
            "Run writing-style before formal delivery.",
            "Run claim-scrutinizer after a full draft exists.",
            "Run red-team to build the strongest opposing case.",
            "Run pre-mortem for failure pathways and leading indicators.",
            "Run boundability for load-bearing moat or NTB claims.",
            "Produce via pattern-docx and run doc-quality-checker.",
        ],
    )
    add_h1(doc, "Release Checklist")
    add_bullets(
        doc,
        [
            "Recommendation appears before supporting detail.",
            "Each gate has evidence, not narrative assertion.",
            "Every NTB has decision impact and kill trigger.",
            "Market section explains why now.",
            "Financial section ties operating drivers to returns.",
            "Risk section includes leading indicators.",
            "Open items are decision-relevant.",
        ],
    )
    add_h1(doc, "Paste-Ready Prompt")
    add_callout(
        doc,
        "Prompt",
        "Write the full Pattern IC memo for [Company]. Decision context: [first look / pre-LOI / final IC]. Deal type: [buyout / acquisition / minority / strategic]. Entry valuation: [details]. Use the full IC memo workflow: mckinsey-consultant, investment-evaluation framework, analytical-operating-system, evidence base, 4-7 NTBs, 10-section memo, returns bridge, scenarios, risks, kill criteria, writing-style, claim-scrutinizer, red-team, pre-mortem, boundability where relevant, pattern-docx, and doc-quality-checker.",
        fill="F5F8FF",
    )
    save(doc, DOCS / "IC_Memo_Gold_Standard_Guide.docx")


def build_competitive_assessment_doc() -> None:
    doc = setup_doc(
        "Competitive Assessment Gold Standard Guide",
        f"How to create the most thorough competitive and moat assessment | {DATE_LABEL}",
    )
    add_callout(
        doc,
        "BLUF",
        "A competitive assessment must answer what position the company can defend, against whom, for how long, and why. It must separate current advantage from durable moat.",
    )
    add_h1(doc, "Decision Standard")
    add_bullets(
        doc,
        [
            "Define the competitive arena before listing competitors.",
            "Explain customer choice, switching threshold, and procurement path.",
            "Include substitutes and platform-native threats.",
            "Prove each moat claim with mechanism, metric, and evidence.",
            "State replicability horizon and erosion vector.",
        ],
    )
    add_h1(doc, "Canonical Assessment Architecture")
    add_table(doc, ["Section", "Minimum standard"], COMPETITIVE_ARCH, [3300, 7500])
    add_h1(doc, "Moat Type Evidence")
    add_table(
        doc,
        ["Moat type", "What must be proven"],
        [
            ["Network effects", "Value per user increases as more users or counterparties join."],
            ["Switching costs", "Customer faces measurable time, cost, risk, or disruption to leave."],
            ["Scale economies", "A named cost line declines structurally with volume or density."],
            ["Proprietary assets", "Data, IP, licenses, brand, or relationships cannot be replicated quickly."],
            ["Efficient scale", "Market is too small or local for rational duplicate infrastructure."],
        ],
        [2800, 8000],
    )
    add_h1(doc, "Required Moat Table")
    add_table(
        doc,
        ["Field", "Standard"],
        [
            ["Moat claim", "Specific claimed advantage, not generic strength."],
            ["Type", "One moat type, assessed independently."],
            ["Evidence", "Named evidence and source quality."],
            ["Strength 1-5", "Rating tied to measurable displacement resistance."],
            ["Replicability horizon", "Time and capital required for a competitor to copy or route around it."],
            ["Erosion vector", "Most likely way the advantage weakens."],
            ["Verdict", "STRONG, MODERATE, WEAK, NOMINAL, or UNPROVEN."],
        ],
        [2800, 8000],
    )
    add_h1(doc, "Quality Gates")
    add_bullets(
        doc,
        [
            "Competitors and substitutes are named.",
            "Every moat claim has a mechanism and metric.",
            "Switching costs are quantified in dollars, time, workflow risk, or contract friction.",
            "Network effects include a feedback loop.",
            "Scale claims name the cost line that improves.",
            "Verdict includes durability and erosion vector.",
        ],
    )
    add_h1(doc, "Paste-Ready Prompt")
    add_callout(
        doc,
        "Prompt",
        "Create a full competitive assessment for [company/category]. Decision: [investment / M&A / product strategy / market entry]. Include direct competitors, substitutes, platform threats, customer buying criteria, competitor map, moat type classification, evidence scorecard, durability and erosion risk, displacement paths, and strategic implications. Use market-research for evidence and competitive-moat-assessment for moat proof. Run writing-style and claim-scrutinizer before DOCX output.",
        fill="F5F8FF",
    )
    save(doc, DOCS / "Competitive_Assessment_Gold_Standard_Guide.docx")


def build_executive_summary_doc() -> None:
    doc = setup_doc(
        "Executive Summary Gold Standard Guide",
        f"How to create a two-page decision-grade executive summary | {DATE_LABEL}",
    )
    add_callout(
        doc,
        "BLUF",
        "The executive summary is the shortest decision-grade version of the argument. It must stand alone for a senior reader and preserve the actual recommendation posture.",
    )
    add_h1(doc, "Canonical Two-Page Spine")
    add_table(doc, ["Section", "What it must answer"], EXEC_SUMMARY_SPINE, [3000, 7800])
    add_h1(doc, "Extraction Before Writing")
    add_table(
        doc,
        ["Section", "Required extraction"],
        [
            ["Company Overview", "Description, customers, scale, stage/status, decision relevance."],
            ["Product Offering", "Workflows, functionality, users, integrations, differentiation."],
            ["Market Dynamic", "Size/growth, timing, demand drivers, competition, headwinds."],
            ["Business Model", "Revenue model, pricing, margins, retention, sales motion, scalability."],
            ["Thesis", "3-5 falsifiable beliefs required for the opportunity to work."],
            ["Open Questions", "Gaps that could change proceed/reprice/pass posture."],
        ],
        [3000, 7800],
    )
    add_h1(doc, "Differentiation Labels")
    add_table(
        doc,
        ["Label", "Definition"],
        [
            ["Structural", "Hard to copy because of data, workflow embed, distribution, scale, switching costs, regulation, or network effects."],
            ["Transient", "Copyable feature, UI, implementation, integration, messaging, or service lead."],
            ["Mixed", "Durable elements exist but execution or current market position still matters."],
            ["Unproven", "Source evidence does not support a durability claim."],
        ],
        [2400, 8400],
    )
    add_h1(doc, "Quality Gates")
    add_bullets(
        doc,
        [
            "The summary can stand alone without the full report.",
            "It preserves the actual recommendation posture.",
            "Every number is sourced or labeled.",
            "Product differentiation is classified.",
            "Thesis bullets are beliefs, not facts.",
            "Open questions are decision-changing.",
        ],
    )
    add_h1(doc, "Paste-Ready Prompt")
    add_callout(
        doc,
        "Prompt",
        "Write a two-page executive summary for [source document / company / market]. Use the canonical six-section spine: Company Overview, Product Offering, Market Dynamic, Business Model, Thesis: What You Need To Believe, and Open Questions. Classify product differentiation as structural, transient, mixed, or unproven. Label facts, estimates, hypotheses, and gaps. Keep it decision-grade and standalone.",
        fill="F5F8FF",
    )
    save(doc, DOCS / "Executive_Summary_Gold_Standard_Guide.docx")


def build_strategic_diligence_doc() -> None:
    doc = setup_doc(
        "Strategic Diligence Gold Standard Guide",
        f"How to create the most thorough NTB and underwriting diligence package | {DATE_LABEL}",
    )
    add_callout(
        doc,
        "BLUF",
        "Strategic diligence is an evidence operating system for a deal: what must be true, what evidence exists, what is missing, and what finding makes us stop.",
    )
    add_h1(doc, "Canonical Diligence Architecture")
    add_table(doc, ["Section", "Minimum standard"], DILIGENCE_ARCH, [3300, 7500])
    add_h1(doc, "Evidence States")
    add_table(
        doc,
        ["State", "Meaning"],
        [
            ["Confirmed", "Direct evidence supports the belief."],
            ["Supported", "Multiple sources support, but not fully proven."],
            ["Mixed", "Evidence conflicts or depends on segment/time period."],
            ["Weak", "Mostly narrative, vendor, or management assertion."],
            ["Missing", "No evidence yet."],
        ],
        [2400, 8400],
    )
    add_h1(doc, "Diligence Plan Fields")
    add_table(
        doc,
        ["Field", "Standard"],
        [
            ["Data request", "Specific data, file, call, benchmark, or analysis required."],
            ["NTB tested", "The load-bearing belief this request confirms or rejects."],
            ["Source owner", "Management, data room, third party, internal model, customer call."],
            ["Format needed", "Workbook, cohort table, contract sample, transcript, benchmark, etc."],
            ["Decision impact", "Proceed, reprice, resolve first, or pass implication."],
            ["Priority", "High only if the answer can change the recommendation."],
        ],
        [2800, 8000],
    )
    add_h1(doc, "Handoff Map")
    add_table(
        doc,
        ["Diligence output", "Downstream use"],
        [
            ["NTB registry", "IC memo investment thesis."],
            ["Evidence state table", "Executive summary and open questions."],
            ["Diligence plan", "Open diligence items."],
            ["Stress tests", "Risks and mitigants."],
            ["Kill triggers", "Recommendation and decision posture."],
            ["Model handoff", "Returns and scenario analysis."],
        ],
        [3300, 7500],
    )
    add_h1(doc, "Quality Gates")
    add_bullets(
        doc,
        [
            "NTBs are not generic risks.",
            "Each NTB has a named evidence state.",
            "Each data request has decision impact.",
            "Kill triggers are specific and measurable.",
            "Stress tests include financial implications.",
            "Compound risks are identified.",
            "Handoff to IC memo and financial model is explicit.",
        ],
    )
    add_h1(doc, "Paste-Ready Prompt")
    add_callout(
        doc,
        "Prompt",
        "Create a full strategic diligence package for [Company]. Decision context: [screen / pre-LOI / IC prep]. Working thesis: [one sentence]. Use ntb-diligence in Full mode with mckinsey-consultant and analytical-operating-system loaded. Build a 4-7 item NTB registry, assign evidence states, create a diligence plan, stress-test every NTB, define kill triggers, map compound risks, and hand off findings to IC memo sections and financial model assumptions.",
        fill="F5F8FF",
    )
    save(doc, DOCS / "Strategic_Diligence_Gold_Standard_Guide.docx")


def main() -> None:
    build_readme_doc()
    build_cheatsheet_doc()
    build_finance_doc()
    build_market_research_doc()
    build_ic_memo_doc()
    build_competitive_assessment_doc()
    build_executive_summary_doc()
    build_strategic_diligence_doc()


if __name__ == "__main__":
    main()
